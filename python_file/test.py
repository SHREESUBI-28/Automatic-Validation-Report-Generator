import pandas as pd
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches

# -----------------------------
# SETUP PROJECT FOLDERS
# -----------------------------
print("Working Directory:", os.getcwd())

os.makedirs("graphs", exist_ok=True)
os.makedirs("reports", exist_ok=True)

# -----------------------------
# AUTO DETECT INPUT FILE
# -----------------------------
input_folder = "input"

files = os.listdir(input_folder)

if len(files) == 0:
    print("No file found in input folder")
    exit()

file_path = os.path.join(input_folder, files[0])

print("Processing:", file_path)

df = pd.read_csv(file_path)

# -----------------------------
# DATA PROCESSING
# -----------------------------
df["Output_Power_kW"] = (df["rb_voltage"] * df["rb_current"]) / 1000

# Safe efficiency calculation
df["Efficiency"] = df.apply(
    lambda x: (x["Output_Power_kW"] / x["Total_kW"] * 100)
    if x["Total_kW"] != 0 else None,
    axis=1
)

# Remove OFF-state samples
df_valid = df[df["rb_current"] > 0].copy()

# -----------------------------
# STATISTICS
# -----------------------------
max_eff = df_valid["Efficiency"].max()
min_eff = df_valid["Efficiency"].min()
avg_eff = df_valid["Efficiency"].mean()

max_power = df_valid["Output_Power_kW"].max()
avg_power = df_valid["Output_Power_kW"].mean()

max_temp = df_valid["IVT_temp"].max()
avg_temp = df_valid["IVT_temp"].mean()

avg_voltage = df_valid["rb_voltage"].mean()

# -----------------------------
# PASS / FAIL LOGIC
# -----------------------------
efficiency_result = "PASS" if avg_eff >= 90 else "FAIL"
temperature_result = "PASS" if max_temp <= 50 else "FAIL"

overall_result = "PASS" if efficiency_result == "PASS" and temperature_result == "PASS" else "FAIL"

# -----------------------------
# PLOTS
# -----------------------------

plt.figure(figsize=(10, 5))
plt.plot(df_valid["Efficiency"])
plt.title("Efficiency vs Sample")
plt.xlabel("Sample")
plt.ylabel("Efficiency (%)")
plt.grid(True)
plt.savefig("graphs/efficiency.png")
plt.close()

plt.figure(figsize=(10, 5))
plt.plot(df_valid["rb_voltage"])
plt.title("Output Voltage vs Sample")
plt.xlabel("Sample")
plt.ylabel("Voltage (V)")
plt.grid(True)
plt.savefig("graphs/output_voltage.png")
plt.close()

plt.figure(figsize=(10, 5))
plt.plot(df_valid["IVT_temp"])
plt.title("Temperature vs Sample")
plt.xlabel("Sample")
plt.ylabel("Temperature (°C)")
plt.grid(True)
plt.savefig("graphs/temperature.png")
plt.close()

plt.figure(figsize=(10, 5))
plt.plot(df_valid["Output_Power_kW"])
plt.title("Output Power vs Sample")
plt.xlabel("Sample")
plt.ylabel("Power (kW)")
plt.grid(True)
plt.savefig("graphs/output_power.png")
plt.close()

# -----------------------------
# SAVE PROCESSED DATA
# -----------------------------
df.to_csv("csv file/processed_rectifier_data.csv", index=False)

print("\nProcessed data saved successfully!")

# -----------------------------
# PRINT SUMMARY
# -----------------------------
print("\n----- TEST EVALUATION -----")
print("Efficiency :", efficiency_result)
print("Temperature :", temperature_result)
print("Overall Result :", overall_result)

print("\n----- VALIDATION SUMMARY -----")
print("Maximum Efficiency :", round(max_eff, 2), "%")
print("Minimum Efficiency :", round(min_eff, 2), "%")
print("Average Efficiency :", round(avg_eff, 2), "%")

print("Maximum Output Power :", round(max_power, 2), "kW")
print("Average Output Power :", round(avg_power, 2), "kW")

print("Maximum Temperature :", round(max_temp, 2), "°C")
print("Average Temperature :", round(avg_temp, 2), "°C")

print("Average Output Voltage :", round(avg_voltage, 2), "V")

# -----------------------------
# REPORT FILE (SUMMARY)
# -----------------------------
with open("reports/validation_summary.txt", "w") as f:

    f.write("RECTIFIER VALIDATION SUMMARY\n")
    f.write("=" * 50 + "\n\n")

    f.write(f"Maximum Efficiency : {max_eff:.2f}%\n")
    f.write(f"Minimum Efficiency : {min_eff:.2f}%\n")
    f.write(f"Average Efficiency : {avg_eff:.2f}%\n\n")

    f.write(f"Maximum Output Power : {max_power:.2f} kW\n")
    f.write(f"Average Output Power : {avg_power:.2f} kW\n\n")

    f.write(f"Maximum Temperature : {max_temp:.2f} °C\n")
    f.write(f"Average Temperature : {avg_temp:.2f} °C\n\n")

    f.write(f"Average Output Voltage : {avg_voltage:.2f} V\n\n")

    f.write("TEST EVALUATION\n")
    f.write("=" * 50 + "\n")
    f.write(f"Efficiency : {efficiency_result}\n")
    f.write(f"Temperature : {temperature_result}\n")
    f.write(f"Overall Result : {overall_result}\n")

print("Validation Summary Generated Successfully")

# -----------------------------
# OBSERVATION FILE
# -----------------------------
with open("reports/observations.txt", "w") as f:

    f.write("OBSERVATIONS\n")
    f.write("=" * 50 + "\n\n")

    f.write(f"Maximum efficiency observed: {max_eff:.2f}%\n")
    f.write(f"Average efficiency observed: {avg_eff:.2f}%\n")
    f.write(f"Maximum output power: {max_power:.2f} kW\n")
    f.write(f"Maximum temperature: {max_temp:.2f} °C\n")

print("Observations Generated Successfully")
# -----------------------------
# WORD REPORT GENERATION
# -----------------------------
document = Document()

# Title
document.add_heading(
    'Rectifier Validation Report',
    level=1
)

# Objective
document.add_heading(
    '1. Objective',
    level=2
)

document.add_paragraph(
    'To evaluate the performance, efficiency and thermal behavior of the rectifier under specified operating conditions.'
)

# Test Setup
document.add_heading(
    '2. Test Setup',
    level=2
)

document.add_paragraph(
    'The rectifier DUT was tested using the validation test setup. Electrical parameters such as voltage, current, power and temperature were monitored and recorded throughout the test duration.'
)

# Validation Summary
document.add_heading(
    '3. Validation Summary',
    level=2
)

document.add_paragraph(
    f"Maximum Efficiency : {max_eff:.2f}%"
)

document.add_paragraph(
    f"Minimum Efficiency : {min_eff:.2f}%"
)

document.add_paragraph(
    f"Average Efficiency : {avg_eff:.2f}%"
)

document.add_paragraph(
    f"Maximum Output Power : {max_power:.2f} kW"
)

document.add_paragraph(
    f"Average Output Power : {avg_power:.2f} kW"
)

document.add_paragraph(
    f"Maximum Temperature : {max_temp:.2f} °C"
)

document.add_paragraph(
    f"Average Temperature : {avg_temp:.2f} °C"
)

document.add_paragraph(
    f"Average Output Voltage : {avg_voltage:.2f} V"
)

# Test Evaluation
document.add_heading(
    '4. Test Evaluation',
    level=2
)

document.add_paragraph(
    f"Efficiency Result : {efficiency_result}"
)

document.add_paragraph(
    f"Temperature Result : {temperature_result}"
)

document.add_paragraph(
    f"Overall Result : {overall_result}"
)

# Observations
document.add_heading(
    '5. Observations',
    level=2
)

if avg_eff >= 90:
    document.add_paragraph(
        f"The rectifier achieved an average efficiency of {avg_eff:.2f}% which meets the acceptance criteria."
    )
else:
    document.add_paragraph(
        f"The rectifier achieved an average efficiency of {avg_eff:.2f}% which is below the required 90% acceptance criteria."
    )

if max_temp <= 50:
    document.add_paragraph(
        f"Thermal performance was satisfactory with a maximum temperature of {max_temp:.2f} °C."
    )
else:
    document.add_paragraph(
        f"Thermal performance was unsatisfactory as the maximum temperature reached {max_temp:.2f} °C."
    )

document.add_paragraph(
    f"The maximum output power delivered during the test was {max_power:.2f} kW."
)

document.add_paragraph(
    f"The average output voltage during the test was {avg_voltage:.2f} V."
)
# Graphs
document.add_page_break()

document.add_heading(
    '6. Graphs',
    level=2
)

efficiency_graph = os.path.abspath("graphs/efficiency.png")
voltage_graph = os.path.abspath("graphs/output_voltage.png")
temperature_graph = os.path.abspath("graphs/temperature.png")
power_graph = os.path.abspath("graphs/output_power.png")

print("Efficiency Graph Path:", efficiency_graph)
print("Voltage Graph Path:", voltage_graph)
print("Temperature Graph Path:", temperature_graph)
print("Power Graph Path:", power_graph)

document.add_paragraph("Efficiency Graph")
document.add_picture(
    efficiency_graph,
    width=Inches(6)
)

document.add_paragraph("Output Voltage Graph")
document.add_picture(
    voltage_graph,
    width=Inches(6)
)

document.add_paragraph("Temperature Graph")
document.add_picture(
    temperature_graph,
    width=Inches(6)
)

document.add_paragraph("Output Power Graph")
document.add_picture(
    power_graph,
    width=Inches(6)
)
# Conclusion
document.add_heading(
    '7. Conclusion',
    level=2
)

if overall_result == "PASS":

    document.add_paragraph(
        f"The rectifier successfully passed the validation test. "
        f"The average efficiency was {avg_eff:.2f}% and the maximum temperature remained within acceptable limits. "
        f"The DUT is considered suitable for operation under the tested conditions."
    )

else:

    document.add_paragraph(
        f"The rectifier failed to meet all validation requirements. "
        f"The average efficiency was {avg_eff:.2f}% which resulted in a FAIL condition. "
        f"Further investigation and optimization are recommended."
    )
# Save Report
document.save(
    "reports/rectifier_report.docx"
)

print("\nWord Report Generated Successfully!")